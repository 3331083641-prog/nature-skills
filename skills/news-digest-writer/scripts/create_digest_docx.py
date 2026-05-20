#!/usr/bin/env python3
"""Create a Word report for the India-ASEAN agriculture monitoring digest."""

from __future__ import annotations

import argparse
import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(run, size: int, bold: bool = False, color: str | None = None) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_paragraph(
    doc: Document,
    text: str,
    size: int = 12,
    bold: bool = False,
    align=None,
    color: str | None = None,
):
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    para.paragraph_format.line_spacing = 1.5
    run = para.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return para


def add_unit_date_line(doc: Document, unit: str, date_short: str):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    tab_stops = para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT)
    unit_run = para.add_run(unit)
    set_run_font(unit_run, size=14, bold=True)
    para.add_run("\t")
    date_run = para.add_run(date_short)
    set_run_font(date_run, size=14, bold=True)
    add_bottom_border(para)
    return para


def set_document_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)


def build_docx(payload: dict, output: Path) -> None:
    doc = Document()
    set_document_margins(doc)

    title = add_paragraph(
        doc,
        "印度、东盟农业政策跟踪",
        size=36,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color="FF0000",
    )
    title.paragraph_format.space_after = Pt(0)

    issue_number = payload.get("issue_number") or "{待确认}"
    add_paragraph(doc, f"总第{issue_number}期", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    date_short = payload["date_short"]
    add_unit_date_line(doc, "华南农业大学经济管理学院", date_short)

    add_paragraph(
        doc,
        f"{payload['date_text']}印度、东盟农业政策动态",
        size=16,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color="08199A",
    )
    add_paragraph(doc, "【本期导读】", size=14, bold=True)

    guide_items = payload.get("guide") or [item["title"] + "。" for item in payload.get("items", [])]
    for guide in guide_items:
        add_paragraph(doc, guide, size=12)

    for item in payload.get("items", []):
        doc.add_page_break()
        add_paragraph(doc, item["title"], size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        body = add_paragraph(doc, item["body"], size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        body.paragraph_format.first_line_indent = Pt(21)
        url = add_paragraph(doc, item["url"], size=12)
        url.paragraph_format.first_line_indent = Pt(21)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def main() -> int:
    parser = argparse.ArgumentParser(description="Create the agriculture monitoring Word digest.")
    parser.add_argument("payload_json", help="Path to report payload JSON")
    parser.add_argument("output_docx", help="Output .docx path")
    args = parser.parse_args()

    payload = json.loads(Path(args.payload_json).read_text(encoding="utf-8"))
    build_docx(payload, Path(args.output_docx))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
